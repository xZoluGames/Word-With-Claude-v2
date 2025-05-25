"""
Sistema Avanzado de Marcas de Agua para Documentos - Versión Actualizada
Compatible con diferentes versiones de python-docx
Incluye configuraciones específicas de posición y tamaño
"""
import tempfile
import atexit
from contextlib import contextmanager
from typing import Optional, Dict, Any
import os
from typing import Optional
from PIL import Image, ImageEnhance
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import base64
from utils.logger import get_logger
from utils.cache import cached, image_cache
from config.settings import RESOURCES_PATHS
logger = get_logger('WatermarkManager')
class WatermarkManager:
    """
    Gestor de marcas de agua con cache integrado.
    
    Mejoras implementadas:
    - Cache de imágenes procesadas
    - Logging de operaciones
    - Configuración desde settings
    """
    def __init__(self):
        self.default_opacity = 0.3
        self.default_position = 'header'
        self._temp_files = set()
        
        # Configuración del encabezado
        self.header_config = {
            'width': Cm(20),
            'height': Cm(27.75),
            'h_align': 'center',
            'v_position': Cm(-1.16),
            'behind_text': True
        }
        
        # Configuración de la insignia
        self.logo_config = {
            'width': Cm(4.36),
            'height': Cm(5.33),
            'align': 'center'
        }
        
        # Registrar limpieza al salir
        atexit.register(self._cleanup_all_temps)
        
        logger.info("WatermarkManager inicializado con limpieza automática")

    def _cleanup_all_temps(self):
        """Limpia todos los archivos temporales al salir"""
        for temp_file in self._temp_files:
            try:
                if os.path.exists(temp_file):
                    os.unlink(temp_file)
                    logger.debug(f"Archivo temporal eliminado: {temp_file}")
            except Exception as e:
                logger.debug(f"Error eliminando temporal: {e}")
        self._temp_files.clear()
    
    @contextmanager
    def _temp_image_file(self, image_data: bytes, suffix: str = '.png'):
        """Context manager para archivos temporales seguros"""
        temp_file = None
        try:
            with tempfile.NamedTemporaryFile(
                mode='wb', suffix=suffix, delete=False
            ) as temp_file:
                temp_file.write(image_data)
                temp_path = temp_file.name
                self._temp_files.add(temp_path)
            
            yield temp_path
            
        finally:
            # Limpiar inmediatamente si es posible
            if temp_file and os.path.exists(temp_file.name):
                try:
                    os.unlink(temp_file.name)
                    self._temp_files.discard(temp_file.name)
                except Exception:
                    pass  # Se limpiará en atexit

    @cached(ttl=86400, key_prefix="watermark_v2")
    def process_image_for_watermark(self, image_path: str, opacity: float = None, 
                                  config: Dict[str, Any] = None) -> Optional[bytes]:
        """
        Procesa imagen con configuración mejorada y validación.
        """
        logger.debug(f"Procesando imagen para marca de agua: {image_path}")
        
        if opacity is None:
            opacity = self.default_opacity
        
        # Validar parámetros
        if not 0 <= opacity <= 1:
            logger.warning(f"Opacidad fuera de rango: {opacity}, usando default")
            opacity = self.default_opacity
        
        if not config:
            config = self.header_config
        
        try:
            # Verificar archivo
            if not os.path.exists(image_path):
                logger.error(f"Archivo no encontrado: {image_path}")
                return None
            
            # Verificar tamaño del archivo
            file_size = os.path.getsize(image_path)
            if file_size > 10 * 1024 * 1024:  # 10MB
                logger.warning(f"Imagen muy grande: {file_size / 1024 / 1024:.1f}MB")
            
            # Abrir y procesar imagen
            with Image.open(image_path) as img:
                logger.debug(f"Imagen abierta: {img.size}, modo: {img.mode}")
                
                # Convertir a RGBA si es necesario
                if img.mode != 'RGBA':
                    img = img.convert('RGBA')
                
                # Redimensionar si se especifica en config
                if 'width' in config:
                    # Calcular dimensiones manteniendo proporción
                    width_cm = config['width'] / Cm(1)
                    width_px = int(width_cm * 96 / 2.54)  # 96 DPI
                    
                    ratio = width_px / img.width
                    height_px = int(img.height * ratio)
                    
                    # Limitar tamaño máximo
                    max_dimension = 2000
                    if width_px > max_dimension or height_px > max_dimension:
                        scale = max_dimension / max(width_px, height_px)
                        width_px = int(width_px * scale)
                        height_px = int(height_px * scale)
                    
                    img = img.resize((width_px, height_px), Image.Resampling.LANCZOS)
                    logger.debug(f"Imagen redimensionada a: {width_px}x{height_px}")
                
                # Aplicar transparencia
                if opacity < 1.0:
                    # Crear nueva imagen con transparencia
                    img_with_alpha = img.copy()
                    alpha = img_with_alpha.split()[-1]
                    alpha = ImageEnhance.Brightness(alpha).enhance(opacity)
                    img_with_alpha.putalpha(alpha)
                    img = img_with_alpha
                
                # Optimizar imagen
                buffer = io.BytesIO()
                img.save(buffer, format='PNG', optimize=True)
                buffer.seek(0)
                
                result = buffer.getvalue()
                logger.info(f"Imagen procesada: {len(result) / 1024:.1f}KB")
                
                return result
                
        except Exception as e:
            logger.error(f"Error procesando imagen: {e}", exc_info=True)
            return None
    
    def configurar_imagen_detras_texto(self, picture, config: dict = None):
        """
        Configura una imagen para que aparezca detrás del texto.
        
        Args:
            picture: Objeto imagen de python-docx
            config: Configuración de posición
        """
        logger.debug("Configurando imagen detrás del texto")
        
        if config is None:
            config = self.header_config
        
        try:
            inline = picture._inline
            
            # Crear elemento anchor
            anchor = OxmlElement('wp:anchor')
            
            # Atributos básicos
            anchor.set('distT', '0')
            anchor.set('distB', '0')
            anchor.set('distL', '0')
            anchor.set('distR', '0')
            anchor.set('simplePos', '0')
            anchor.set('relativeHeight', '0')
            anchor.set('behindDoc', '1')  # Detrás del texto
            anchor.set('locked', '0')
            anchor.set('layoutInCell', '1')
            anchor.set('allowOverlap', '1')
            
            # Posición simple (requerida)
            simplePos = OxmlElement('wp:simplePos')
            simplePos.set('x', '0')
            simplePos.set('y', '0')
            anchor.append(simplePos)
            
            # Posición horizontal - Centrada
            positionH = OxmlElement('wp:positionH')
            positionH.set('relativeFrom', 'page')
            alignH = OxmlElement('wp:align')
            alignH.text = 'center'
            positionH.append(alignH)
            anchor.append(positionH)
            
            # Posición vertical
            positionV = OxmlElement('wp:positionV')
            positionV.set('relativeFrom', 'paragraph')
            posOffsetV = OxmlElement('wp:posOffset')
            posOffsetV.text = str(int(config.get('v_position', Cm(-1.5))))
            positionV.append(posOffsetV)
            anchor.append(positionV)
            
            # Tamaño
            extent = OxmlElement('wp:extent')
            extent.set('cx', str(int(config.get('width', Cm(20.96)))))
            extent.set('cy', str(int(config.get('height', Cm(27.68)))))
            anchor.append(extent)
            
            # Efecto visual
            wrapNone = OxmlElement('wp:wrapNone')
            anchor.append(wrapNone)
            
            # Copiar elementos del inline al anchor
            for element in inline:
                if element.tag.endswith(('docPr', 'cNvGraphicFramePr', 'graphic')):
                    anchor.append(element)
            
            # Reemplazar inline con anchor
            inline.getparent().replace(inline, anchor)
            
            logger.debug("Imagen configurada detrás del texto exitosamente")
            
        except Exception as e:
            logger.error(f"Error configurando imagen detrás del texto: {e}", exc_info=True)
            raise
    
    def add_watermark_to_section(self, section, image_path: str, 
                               opacity: float = None, stretch: bool = True,
                               mode: str = 'watermark') -> bool:
        """
        Agrega marca de agua con modo mejorado.
        """
        logger.info(f"Agregando marca de agua - Modo: {mode}, Opacidad: {opacity}")
        
        if opacity is None:
            opacity = self.default_opacity
        
        try:
            # Procesar imagen según modo
            if mode == 'watermark' and opacity < 1.0:
                # Procesar con transparencia
                processed_image = self.process_image_for_watermark(
                    image_path, opacity, self.header_config
                )
                
                if processed_image:
                    # Usar imagen procesada
                    with self._temp_image_file(processed_image) as temp_path:
                        return self._add_image_to_header(section, temp_path, stretch)
                else:
                    logger.warning("Fallo al procesar imagen, usando original")
                    return self._add_image_to_header(section, image_path, stretch)
            else:
                # Modo normal sin transparencia
                return self._add_image_to_header(section, image_path, stretch)
                
        except Exception as e:
            logger.error(f"Error agregando marca de agua: {e}", exc_info=True)
            return False
    def _add_image_to_header(self, section, image_path: str, stretch: bool) -> bool:
        """Método interno para agregar imagen al header"""
        try:
            # Configurar sección
            section.header_distance = Cm(1.25)
            section.footer_distance = Cm(1.25)
            
            # Obtener o crear header
            header = section.header
            if not header.paragraphs:
                header_para = header.add_paragraph()
            else:
                header_para = header.paragraphs[0]
                # Limpiar contenido existente
                for run in header_para.runs:
                    run.clear()
            
            # Configurar alineación
            header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Agregar imagen
            run = header_para.add_run()
            
            if stretch:
                header_pic = run.add_picture(image_path, width=self.header_config['width'])
            else:
                # Mantener proporción original
                header_pic = run.add_picture(image_path)
                
                # Ajustar si es muy grande
                if header_pic.width > self.header_config['width']:
                    ratio = self.header_config['width'] / header_pic.width
                    header_pic.width = self.header_config['width']
                    header_pic.height = int(header_pic.height * ratio)
            
            # Intentar configurar detrás del texto
            try:
                self.configurar_imagen_detras_texto(header_pic, self.header_config)
                logger.info("✅ Imagen configurada detrás del texto")
            except Exception as e:
                logger.warning(f"⚠️ No se pudo configurar detrás del texto: {e}")
            
            return True
            
        except Exception as e:
            logger.error(f"Error agregando imagen al header: {e}")
            return False
    def add_logo_to_first_page(self, doc, logo_path: str) -> bool:
        """
        Agrega logo/insignia a la primera página.
        
        Args:
            doc: Documento
            logo_path: Ruta del logo
            
        Returns:
            bool: True si se agregó exitosamente
        """
        logger.info(f"Agregando logo a primera página: {logo_path}")
        
        if not logo_path or not os.path.exists(logo_path):
            logger.error(f"Logo no encontrado: {logo_path}")
            return False
        
        try:
            # Agregar párrafo para el logo
            logo_para = doc.add_paragraph()
            logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Usar imagen del cache si es posible
            logo_run = logo_para.add_run()
            logo_run.add_picture(logo_path, height=self.logo_config['height'])
            
            # Espacio después del logo
            doc.add_paragraph()
            
            logger.info("✅ Logo agregado a primera página")
            return True
            
        except Exception as e:
            logger.error(f"Error agregando logo: {e}", exc_info=True)
            return False
    
    def clear_cache(self):
        """Limpia el cache de imágenes procesadas."""
        logger.info("Limpiando cache de marcas de agua")
        
        # Invalidar cache de funciones decoradas
        self.process_image_for_watermark.invalidate_cache()
        
        # Limpiar cache de imágenes
        if hasattr(image_cache, 'clear_image_cache'):
            image_cache.clear_image_cache()
        
        logger.info("Cache de marcas de agua limpiado")
    
    def _add_watermark_alternative(self, paragraph, image_path, opacity, stretch):
        """Método alternativo para agregar marca de agua usando XML directo"""
        try:
            # Procesar imagen primero
            if stretch:
                # Convertir cm a pulgadas para procesamiento
                width_inches = self.header_config['width'] / Cm(1) / 2.54
                processed_image = self.process_image_for_watermark(image_path, opacity, width_inches)
            else:
                processed_image = self.process_image_for_watermark(image_path, opacity, 5)
            
            if not processed_image:
                return False
            
            # Convertir a base64
            image_base64 = base64.b64encode(processed_image).decode('utf-8')
            
            # Crear elemento run
            run = paragraph.add_run()
            r = run._r
            
            # Crear estructura pict
            pict = OxmlElement('w:pict')
            
            # Crear shape con configuración específica
            shape = OxmlElement('v:shape')
            shape.set('id', '_x0000_i1025')
            shape.set('type', '#_x0000_t75')
            
            # Convertir dimensiones a puntos para el estilo
            width_pt = int(self.header_config['width'] / Cm(1) * 28.35)
            height_pt = int(self.header_config['height'] / Cm(1) * 28.35)
            
            shape.set('style', f'width:{width_pt}pt;height:{height_pt}pt;position:absolute;z-index:-251658752')
            
            # Crear imagedata
            imagedata = OxmlElement('v:imagedata')
            imagedata.set(qn('r:id'), 'rId1')
            imagedata.set('o:title', 'Watermark')
            
            # Agregar imagedata a shape
            shape.append(imagedata)
            
            # Agregar shape a pict
            pict.append(shape)
            
            # Agregar pict a run
            r.append(pict)
            
            # Intentar agregar la relación de imagen
            try:
                # Este es un método simplificado, puede necesitar ajustes
                document = paragraph.part
                image_part = document.new_image_part(image_path)
                imagedata.set(qn('r:id'), document.relate_to(image_part, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/image'))
            except:
                pass
            
            return True
            
        except Exception as e:
            print(f"Error en método alternativo: {e}")
            return False
    
    def add_simple_header_image(self, section, image_path, width_inches=None):
        """Método simple para agregar imagen al encabezado con dimensiones correctas"""
        try:
            header = section.header
            
            # Configurar márgenes de sección
            section.header_distance = Cm(1.25)
            section.footer_distance = Cm(1.25)
            
            # Asegurar que hay un párrafo
            if not header.paragraphs:
                p = header.add_paragraph()
            else:
                p = header.paragraphs[0]
            
            # Centrar el párrafo
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Agregar la imagen con configuración específica
            run = p.add_run()
            
            # Usar ancho de configuración si no se especifica
            if width_inches is None:
                picture = run.add_picture(image_path, width=self.header_config['width'])
            else:
                picture = run.add_picture(image_path, width=Inches(width_inches))
            
            return True
            
        except Exception as e:
            print(f"Error agregando imagen simple al header: {e}")
            return False
    
    def configure_document_headers(self, doc, header_image_path, logo_image_path=None):
        """Configura encabezados del documento completo según especificaciones"""
        try:
            # Configurar primera sección
            section = doc.sections[0]
            
            # IMPORTANTE: Primera página diferente
            section.different_first_page_header_footer = True
            
            # Configurar márgenes
            section.top_margin = Cm(2.5)
            section.bottom_margin = Cm(2.5)
            section.left_margin = Cm(3)
            section.right_margin = Cm(3)
            section.header_distance = Cm(1.25)
            section.footer_distance = Cm(1.25)
            
            # Agregar encabezado a páginas 2+
            if header_image_path and os.path.exists(header_image_path):
                self.add_watermark_to_section(section, header_image_path)
            
            print("✅ Configuración de encabezados completada")
            return True
            
        except Exception as e:
            print(f"Error configurando encabezados del documento: {e}")
            return False