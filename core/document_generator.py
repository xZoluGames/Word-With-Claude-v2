
from utils.logger import get_logger

logger = get_logger("document_generator")

"""
Generador de documentos Word - Versión Corregida con Encabezados como Marca de Agua y sangría APA perfecta
"""
from concurrent.futures import ThreadPoolExecutor
from functools import lru_cache
import tempfile
from contextlib import contextmanager
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_BREAK
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.section import WD_SECTION, WD_ORIENTATION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from modules.watermark import WatermarkManager
import threading
import os
from datetime import datetime
from tkinter import filedialog, messagebox
import re

class DocumentGenerator:
    def __init__(self):
        self.formato_config = {
            'fuente_texto': 'Times New Roman',
            'tamaño_texto': 12,
            'fuente_titulo': 'Times New Roman', 
            'tamaño_titulo': 14,
            'interlineado': 2.0,
            'margen': 2.54,
            'justificado': True,
            'sangria': True
        }
        self.watermark_manager = WatermarkManager()
        self._executor = ThreadPoolExecutor(max_workers=2)
        self._temp_files = []

    def __del__(self):
        """Limpieza de recursos"""
        self._executor.shutdown(wait=False)
        self._cleanup_temp_files()

    def _cleanup_temp_files(self):
        """Limpia archivos temporales"""
        for temp_file in self._temp_files:
            try:
                if os.path.exists(temp_file):
                    os.unlink(temp_file)
            except Exception as e:
                logger.debug(f"Error eliminando archivo temporal: {e}")
        self._temp_files.clear()

    @contextmanager
    def _temp_document(self):
        """Context manager para documento temporal"""
        doc = Document()
        try:
            yield doc
        finally:
            # Limpieza si es necesaria
            pass

    @lru_cache(maxsize=32)
    def normalizar_parrafos(self, contenido):
        """
        Normaliza párrafos con cache para contenido repetido.
        Mejorado para manejar diferentes formatos de entrada.
        """
        if not contenido:
            return ""
        
        texto = contenido.strip()
        
        # Si ya tiene dobles saltos, validar formato
        if '\n\n' in texto:
            # Limpiar saltos excesivos (más de 2)
            texto = re.sub(r'\n{3,}', '\n\n', texto)
            return texto
        
        # Detectar si es una lista
        lineas = texto.split('\n')
        es_lista = any(
            linea.strip().startswith(('•', '-', '*', '1.', '2.', '3.'))
            for linea in lineas if linea.strip()
        )
        
        if es_lista:
            # Mantener formato de lista pero asegurar espaciado
            return '\n'.join(linea for linea in lineas if linea.strip())
        
        # Para texto normal, convertir saltos simples en párrafos
        parrafos = []
        parrafo_actual = []
        
        for linea in lineas:
            linea = linea.strip()
            if not linea:
                if parrafo_actual:
                    parrafos.append(' '.join(parrafo_actual))
                    parrafo_actual = []
            else:
                parrafo_actual.append(linea)
        
        if parrafo_actual:
            parrafos.append(' '.join(parrafo_actual))
        
        return '\n\n'.join(parrafos)

    def generar_documento_async(self, app_instance):
        """Versión mejorada con mejor manejo de errores y progreso"""
        def _progress_callback(valor, mensaje=""):
            """Actualiza progreso de forma segura"""
            try:
                app_instance.root.after(0, lambda: app_instance.progress.set(valor))
                if mensaje and hasattr(app_instance, 'status_label'):
                    app_instance.root.after(0, 
                        lambda: app_instance.status_label.configure(text=mensaje))
            except Exception as e:
                logger.debug(f"Error actualizando progreso: {e}")

        def generar():
            try:
                logger.info("Iniciando generación de documento")
                _progress_callback(0, "Iniciando generación...")
                
                # Validar antes de generar
                if not app_instance.validator.validar_proyecto(app_instance):
                    if not messagebox.askyesno("⚠️ Validación", 
                        "El proyecto tiene errores. ¿Generar de todos modos?"):
                        _progress_callback(0, "Generación cancelada")
                        return
                
                with self._temp_document() as doc:
                    # Configuración inicial
                    _progress_callback(0.1, "Configurando documento...")
                    self.configurar_documento_completo(doc, app_instance)
                    
                    # Portada
                    if app_instance.incluir_portada.get():
                        _progress_callback(0.2, "Creando portada...")
                        self.crear_portada_profesional(doc, app_instance)
                    
                    # Agradecimientos
                    if app_instance.incluir_agradecimientos.get():
                        _progress_callback(0.3, "Agregando agradecimientos...")
                        contenido_agradecimientos = self._get_agradecimientos_default()
                        self.crear_seccion_profesional(
                            doc, "AGRADECIMIENTOS", contenido_agradecimientos, 
                            app_instance, nivel=1, aplicar_sangria_parrafos=False
                        )
                    
                    # Resumen
                    if self._tiene_resumen(app_instance):
                        _progress_callback(0.4, "Procesando resumen...")
                        contenido_resumen = app_instance.content_texts['resumen'].get("1.0", "end")
                        contenido_resumen = self.normalizar_parrafos(contenido_resumen)
                        self.crear_seccion_profesional(
                            doc, "RESUMEN", contenido_resumen, app_instance, 
                            nivel=1, aplicar_sangria_parrafos=False
                        )
                    
                    # Índice
                    if app_instance.incluir_indice.get():
                        _progress_callback(0.5, "Generando índice...")
                        self.crear_indice_profesional(doc, app_instance)
                    
                    # Contenido principal
                    _progress_callback(0.6, "Procesando contenido principal...")
                    self.crear_contenido_dinamico_mejorado(doc, app_instance)
                    
                    # Referencias
                    _progress_callback(0.8, "Agregando referencias...")
                    self.crear_referencias_profesionales(doc, app_instance)
                    
                    # Guardar documento
                    _progress_callback(0.9, "Guardando documento...")
                    filename = filedialog.asksaveasfilename(
                        defaultextension=".docx",
                        filetypes=[("Word documents", "*.docx")],
                        title="Guardar Proyecto Académico Profesional",
                        initialfile=self._generar_nombre_archivo(app_instance)
                    )
                    
                    if filename:
                        doc.save(filename)
                        _progress_callback(1.0, "¡Documento generado exitosamente!")
                        
                        # Limpiar archivos temporales
                        self._cleanup_temp_files()
                        
                        self.mostrar_mensaje_exito(filename, app_instance)
                        
                        # Abrir documento si el usuario lo desea
                        if messagebox.askyesno("📄 Abrir Documento", 
                            "¿Deseas abrir el documento generado?"):
                            self._abrir_documento(filename)
                    else:
                        _progress_callback(0, "Generación cancelada")
                
            except Exception as e:
                logger.error(f"Error generando documento: {e}", exc_info=True)
                _progress_callback(0, "Error en la generación")
                messagebox.showerror("❌ Error", 
                    f"Error al generar documento:\n{str(e)}\n\n"
                    "Revisa el log para más detalles.")
        
        # Ejecutar en thread
        thread = threading.Thread(target=generar, daemon=True)
        thread.start()
    
    def configurar_documento_completo(self, doc, app_instance):
        for section in doc.sections:
            section.top_margin = Inches(app_instance.formato_config['margen'] / 2.54)
            section.bottom_margin = Inches(app_instance.formato_config['margen'] / 2.54)
            section.left_margin = Inches(app_instance.formato_config['margen'] / 2.54)
            section.right_margin = Inches(app_instance.formato_config['margen'] / 2.54)
            self.configurar_encabezado_marca_agua(section, app_instance)
        self.configurar_estilos_profesionales(doc, app_instance)
    
    def configurar_encabezado_marca_agua(self, section, app_instance):
        try:
            section.different_first_page_header_footer = True
            section.top_margin = Cm(2.5)
            section.bottom_margin = Cm(2.5)
            section.left_margin = Cm(3)
            section.right_margin = Cm(3)
            section.header_distance = Cm(1.25)
            section.footer_distance = Cm(1.25)
            ruta_encabezado = self.obtener_ruta_imagen("encabezado", app_instance)
            if ruta_encabezado and os.path.exists(ruta_encabezado):
                opacity = getattr(app_instance, 'watermark_opacity', 0.3)
                stretch = getattr(app_instance, 'watermark_stretch', True)
                mode = getattr(app_instance, 'watermark_mode', 'watermark')
                header = section.header
                for para in header.paragraphs:
                    p = para._element
                    p.getparent().remove(p)
                header_para = header.add_paragraph()
                header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = header_para.add_run()
                if mode == 'watermark' and hasattr(self, 'watermark_manager'):
                    try:
                        header_pic = run.add_picture(ruta_encabezado, width=self.watermark_manager.header_config['width'])
                        self.watermark_manager.configurar_imagen_detras_texto(header_pic, self.watermark_manager.header_config)
                    except Exception:
                        self.watermark_manager.add_simple_header_image(section, ruta_encabezado)
                else:
                    run.add_picture(ruta_encabezado, width=Cm(20.96))
            else:
                self._configurar_encabezado_simple(section, app_instance)
        except Exception:
            self._configurar_encabezado_simple(section, app_instance)
    
    def _configurar_encabezado_simple(self, section, app_instance):
        try:
            header = section.header
            for para in header.paragraphs:
                p = para._element
                p.getparent().remove(p)
            p = header.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            institucion = app_instance.proyecto_data.get('institucion', None)
            if institucion and hasattr(institucion, 'get'):
                texto = institucion.get() or "INSTITUCIÓN EDUCATIVA"
            else:
                texto = "INSTITUCIÓN EDUCATIVA"
            run = p.add_run(texto.upper())
            run.font.name = 'Times New Roman'
            run.font.size = Pt(14)
            run.font.bold = True
        except Exception:
            pass
    
    def configurar_estilos_profesionales(self, doc, app_instance):
        style = doc.styles['Normal']
        style.font.name = app_instance.formato_config['fuente_texto']
        style.font.size = Pt(app_instance.formato_config['tamaño_texto'])
        style.font.color.rgb = RGBColor(0, 0, 0)
        if app_instance.formato_config['interlineado'] == 1.0:
            style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        elif app_instance.formato_config['interlineado'] == 1.5:
            style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        else:
            style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        if app_instance.formato_config['justificado']:
            style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        style.paragraph_format.space_after = Pt(0)
        try:
            body_style = doc.styles.add_style('BodyTextIndent', WD_STYLE_TYPE.PARAGRAPH)
            body_style.base_style = doc.styles['Normal']
            body_style.font.name = app_instance.formato_config['fuente_texto']
            body_style.font.size = Pt(app_instance.formato_config['tamaño_texto'])
            if app_instance.formato_config['sangria']:
                body_style.paragraph_format.first_line_indent = Inches(0)
        except:
            if 'BodyTextIndent' in doc.styles:
                body_style = doc.styles['BodyTextIndent']
                body_style.font.name = app_instance.formato_config['fuente_texto']
                body_style.font.size = Pt(app_instance.formato_config['tamaño_texto'])
                if app_instance.formato_config['sangria']:
                    body_style.paragraph_format.first_line_indent = Inches(0.5)
        for i in range(1, 7):
            heading_name = f'Heading {i}'
            if heading_name in doc.styles:
                heading_style = doc.styles[heading_name]
            else:
                try:
                    heading_style = doc.styles.add_style(heading_name, WD_STYLE_TYPE.PARAGRAPH)
                except:
                    continue
            heading_style.font.name = app_instance.formato_config['fuente_titulo']
            heading_style.font.size = Pt(app_instance.formato_config['tamaño_titulo'])
            heading_style.paragraph_format.page_break_before = True
            heading_style.font.bold = True
            heading_style.font.color.rgb = RGBColor(0, 0, 0)
            heading_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            heading_style.paragraph_format.space_before = Pt(12)
            heading_style.paragraph_format.space_after = Pt(12)
            heading_style.paragraph_format.keep_with_next = True
            heading_style.paragraph_format.outline_level = i - 1
            heading_style.paragraph_format.first_line_indent = Inches(0)
    
    def crear_portada_profesional(self, doc, app_instance):
        ruta_insignia = self.obtener_ruta_imagen("insignia", app_instance)
        if ruta_insignia and os.path.exists(ruta_insignia):
            try:
                if hasattr(self, 'watermark_manager'):
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.paragraph_format.first_line_indent = Inches(0)
                    run = p.add_run()
                    run.add_picture(ruta_insignia, height=self.watermark_manager.logo_config['height'])
                else:
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run()
                    run.add_picture(ruta_insignia, width=Inches(1.5))
            except Exception:
                pass
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Inches(0)
        run = p.add_run(app_instance.proyecto_data['institucion'].get().upper())
        run.bold = True
        run.font.name = app_instance.formato_config['fuente_titulo']
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor(0, 0, 0)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Inches(0)
        run = p.add_run(f'"{app_instance.proyecto_data["titulo"].get()}"')
        run.bold = True
        run.font.name = app_instance.formato_config['fuente_titulo']
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor(0, 0, 0)
        info_fields = [
            ('ciclo', 'Ciclo'),
            ('curso', 'Curso'), 
            ('enfasis', 'Énfasis'),
            ('area', 'Área de Desarrollo'),
            ('categoria', 'Categoría'),
            ('director', 'Director'),
            ('responsable', 'Responsable')
        ]
        for field, label in info_fields:
            if field in app_instance.proyecto_data and app_instance.proyecto_data[field].get().strip():
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                label_run = p.add_run(f"{label}: ")
                label_run.bold = True
                label_run.font.name = app_instance.formato_config['fuente_texto']
                label_run.font.size = Pt(14)
                label_run.font.color.rgb = RGBColor(0, 0, 0)
                valor_original = app_instance.proyecto_data[field].get()
                if field == 'responsable' and ',' in valor_original:
                    responsables = [resp.strip() for resp in valor_original.split(',')]
                    if len(responsables) == 1:
                        valor_formateado = responsables[0]
                    elif len(responsables) == 2:
                        valor_formateado = f"{responsables[0]} y {responsables[1]}"
                    else:
                        todos_menos_ultimo = ", ".join(responsables[:-1])
                        valor_formateado = f"{todos_menos_ultimo} y {responsables[-1]}"
                    value_run = p.add_run(valor_formateado)
                else:
                    value_run = p.add_run(valor_original)
                value_run.font.name = app_instance.formato_config['fuente_texto']
                value_run.font.size = Pt(12)
                value_run.font.color.rgb = RGBColor(0, 0, 0)
        if app_instance.proyecto_data['estudiantes'].get():
            self._agregar_lista_personas(doc, "Estudiantes", 
                                    app_instance.proyecto_data['estudiantes'].get(), 
                                    app_instance, alineacion='izquierda')
        if app_instance.proyecto_data['tutores'].get():
            self._agregar_lista_personas(doc, "Tutores", 
                                    app_instance.proyecto_data['tutores'].get(), 
                                    app_instance, alineacion='izquierda')
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        year_label = p.add_run("Año: ")
        year_label.bold = True
        year_label.font.name = app_instance.formato_config['fuente_texto']
        year_label.font.size = Pt(14)
        year_label.font.color.rgb = RGBColor(0, 0, 0)
        year_value = p.add_run(str(datetime.now().year))
        year_value.font.name = app_instance.formato_config['fuente_texto']
        year_value.font.size = Pt(12)
        year_value.font.color.rgb = RGBColor(0, 0, 0)
        doc.add_page_break()

    def _agregar_lista_personas(self, doc, titulo, personas_str, app_instance, alineacion='centro'):
        p = doc.add_paragraph()
        if alineacion == 'izquierda':
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = p.add_run(f"{titulo}: ")
        title_run.bold = True
        title_run.font.name = app_instance.formato_config['fuente_texto']
        title_run.font.size = Pt(14)
        title_run.font.color.rgb = RGBColor(0, 0, 0)
        personas = []
        for persona in personas_str.split(','):
            persona_limpia = persona.strip()
            if persona_limpia:
                personas.append(persona_limpia)
        if len(personas) == 0:
            personas_run = p.add_run("")
        elif len(personas) == 1:
            personas_run = p.add_run(personas[0])
        elif len(personas) == 2:
            personas_run = p.add_run(f"{personas[0]} y {personas[1]}")
        else:
            todos_menos_ultimo = ", ".join(personas[:-1])
            personas_run = p.add_run(f"{todos_menos_ultimo} y {personas[-1]}")
        personas_run.font.name = app_instance.formato_config['fuente_texto']
        personas_run.font.size = Pt(12)
        personas_run.font.color.rgb = RGBColor(0, 0, 0)
    
    def crear_indice_profesional(self, doc, app_instance):
        p = doc.add_heading('ÍNDICE', level=1)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Inches(0)
        doc.add_paragraph()
        instrucciones = """INSTRUCCIONES PARA GENERAR ÍNDICE AUTOMÁTICO:

    1. En Word, ir a la pestaña "Referencias"
    2. Hacer clic en "Tabla de contenido"  
    3. Seleccionar el estilo deseado
    4. El índice se generará automáticamente

    NOTA: Todos los títulos están configurados con niveles de esquema para facilitar la generación automática."""
        for linea in instrucciones.split('\n'):
            p = doc.add_paragraph(linea)
            p.paragraph_format.first_line_indent = Inches(0)
        doc.add_paragraph()
        p = doc.add_heading('TABLA DE ILUSTRACIONES', level=2)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Inches(0)
        p = doc.add_paragraph("(Agregar manualmente si hay figuras, tablas o gráficos)")
        p.paragraph_format.first_line_indent = Inches(0)
        doc.add_page_break()
    
    def crear_contenido_dinamico_mejorado(self, doc, app_instance):
        capitulo_num = 0
        secciones_sin_sangria = [
            "RESUMEN", "PALABRAS CLAVE", "AGRADECIMIENTOS",
            "ÍNDICE", "TABLA DE ILUSTRACIONES", "REFERENCIAS"
        ]
        for seccion_id in app_instance.secciones_activas:
            if seccion_id in app_instance.secciones_disponibles:
                seccion = app_instance.secciones_disponibles[seccion_id]
                if seccion['capitulo']:
                    capitulo_num += 1
                    titulo = seccion['titulo']
                    titulo_limpio = re.sub(r'[^\w\s-]', '', titulo).strip()
                    p = doc.add_heading(titulo_limpio, level=1)
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.paragraph_format.first_line_indent = Inches(0)
                    if capitulo_num > 1:
                        doc.add_page_break()
                    else:
                        doc.add_paragraph()
                else:
                    if seccion_id in app_instance.content_texts:
                        raw_content = app_instance.content_texts[seccion_id].get("1.0", "end")
                        contenido = self.normalizar_parrafos(raw_content)
                        if contenido:
                            titulo = seccion['titulo']
                            titulo_limpio = re.sub(r'[^\w\s-]', '', titulo).strip()
                            aplicar_sangria = not (titulo_limpio.upper() in secciones_sin_sangria)
                            self.crear_seccion_profesional(
                                doc, titulo_limpio.upper(), contenido, app_instance, nivel=2,
                                aplicar_sangria_parrafos=aplicar_sangria
                            )

    def crear_seccion_profesional(self, doc, titulo, contenido, app_instance, nivel=1, aplicar_sangria_parrafos=True):
        p = doc.add_heading(titulo, level=nivel)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Inches(0)
        contenido_procesado = self.procesar_citas_mejorado(contenido.strip(), app_instance)
        parrafos = contenido_procesado.split('\n\n')
        for i, parrafo in enumerate(parrafos):
            texto = parrafo.strip()
            if texto:
                p = doc.add_paragraph(texto)
                aplicar_sangria = aplicar_sangria_parrafos
                if len(texto.split()) > 40:
                    aplicar_sangria = False
                    p.paragraph_format.left_indent = Inches(0.5)
                    p.paragraph_format.right_indent = Inches(0.5)
                    p.paragraph_format.first_line_indent = Inches(0)
                elif texto.startswith('\t') or texto.startswith('     '):
                    aplicar_sangria = False
                    p.paragraph_format.left_indent = Inches(0.5)
                    p.paragraph_format.right_indent = Inches(0.5)
                    p.paragraph_format.first_line_indent = Inches(0)
                elif any(texto.startswith(marca) for marca in [
                    '•', '-', '*', '1.', '2.', '3.', '4.', '5.', '6.', '7.', '8.', '9.', '10.'
                ]):
                    aplicar_sangria = False
                    p.paragraph_format.left_indent = Inches(0.5)
                    p.paragraph_format.first_line_indent = Inches(0)
                if aplicar_sangria and app_instance.formato_config.get('sangria', True):
                    p.paragraph_format.first_line_indent = Inches(0.5)
                else:
                    p.paragraph_format.first_line_indent = Inches(0)
                p.style = doc.styles['Normal']
        doc.add_paragraph()
    
    def procesar_citas_mejorado(self, texto, app_instance):
        if hasattr(app_instance, 'citation_processor'):
            return app_instance.citation_processor.procesar_citas_avanzado(texto)
        def reemplazar_cita(match):
            cita_completa = match.group(0)
            contenido = cita_completa[6:-1]
            partes = contenido.split(':')
            if len(partes) >= 3:
                tipo, autor, año = partes[0], partes[1], partes[2]
                pagina = partes[3] if len(partes) > 3 else None
                if tipo == 'textual':
                    if pagina:
                        return f" ({autor}, {año}, p. {pagina})"
                    else:
                        return f" ({autor}, {año})"
                elif tipo == 'parafraseo':
                    return f" ({autor}, {año})"
                elif tipo == 'larga':
                    if pagina:
                        return f"\n\n     ({autor}, {año}, p. {pagina})\n\n"
                    else:
                        return f"\n\n     ({autor}, {año})\n\n"
                elif tipo == 'web':
                    return f" ({autor}, {año})"
                elif tipo == 'multiple':
                    return f" ({autor}, {año})"
                else:
                    return f" ({autor}, {año})"
            return cita_completa
        texto_procesado = re.sub(r'\[CITA:[^\]]+\]', reemplazar_cita, texto)
        return texto_procesado
    
    def crear_referencias_profesionales(self, doc, app_instance):
        if not app_instance.referencias:
            return
        p = doc.add_heading('REFERENCIAS', level=1)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()
        referencias_ordenadas = sorted(app_instance.referencias, 
                                    key=lambda x: x['autor'].split(',')[0].strip())
        for ref in referencias_ordenadas:
            ref_text = self._formatear_referencia_apa(ref)
            p = doc.add_paragraph(ref_text)
            p.paragraph_format.first_line_indent = Inches(-0.5)
            p.paragraph_format.left_indent = Inches(0.5)
            p.style = doc.styles['Normal']

    def _formatear_referencia_apa(self, ref):
        tipo = ref.get('tipo', 'Libro')
        autor = ref.get('autor', '')
        año = ref.get('año', '')
        titulo = ref.get('titulo', '')
        fuente = ref.get('fuente', '')
        if tipo == 'Libro':
            return f"{autor} ({año}). {titulo}. {fuente}."
        elif tipo == 'Artículo':
            return f"{autor} ({año}). {titulo}. {fuente}."
        elif tipo == 'Web':
            return f"{autor} ({año}). {titulo}. Recuperado de {fuente}"
        elif tipo == 'Tesis':
            return f"{autor} ({año}). {titulo} [Tesis]. {fuente}."
        else:
            return f"{autor} ({año}). {titulo}. {fuente}."
    
    def obtener_ruta_imagen(self, tipo, app_instance):
        if tipo == "encabezado":
            return (getattr(app_instance, 'encabezado_personalizado', None) or 
                   getattr(app_instance, 'ruta_encabezado', None))
        elif tipo == "insignia":
            return (getattr(app_instance, 'insignia_personalizada', None) or 
                   getattr(app_instance, 'ruta_insignia', None))
        return None
    def _tiene_resumen(self, app_instance):
        """Verifica si existe resumen con contenido"""
        return ('resumen' in app_instance.secciones_activas and 
                'resumen' in app_instance.content_texts and
                app_instance.content_texts['resumen'].get("1.0", "end").strip())

    def _get_agradecimientos_default(self):
        """Obtiene texto por defecto para agradecimientos"""
        return ("A mis tutores, por su invaluable guía y apoyo constante.\n\n"
                "A mi familia, por su comprensión y motivación.\n\n"
                "A la institución, por brindarme las herramientas necesarias "
                "para desarrollar este proyecto.")

    def _generar_nombre_archivo(self, app_instance):
        """Genera nombre de archivo sugerido"""
        titulo = app_instance.proyecto_data.get('titulo', {})
        if hasattr(titulo, 'get'):
            titulo = titulo.get()
        
        if titulo:
            # Limpiar título para nombre de archivo
            nombre = re.sub(r'[^\w\s-]', '', titulo)
            nombre = re.sub(r'[-\s]+', '_', nombre)
            return f"{nombre[:50]}.docx"
        
        return f"proyecto_academico_{datetime.now().strftime('%Y%m%d')}.docx"

    def _abrir_documento(self, filepath):
        """Abre el documento en el programa predeterminado"""
        try:
            if sys.platform.startswith('darwin'):
                os.system(f'open "{filepath}"')
            elif sys.platform.startswith('win'):
                os.startfile(filepath)
            else:  # linux
                os.system(f'xdg-open "{filepath}"')
        except Exception as e:
            logger.error(f"Error abriendo documento: {e}")
    def mostrar_mensaje_exito(self, filename, app_instance):
        app_instance.validation_text.delete("1.0", "end")
        app_instance.validation_text.insert("1.0", 
            f"🎉 ¡DOCUMENTO PROFESIONAL GENERADO!\n\n"
            f"📄 Archivo: {os.path.basename(filename)}\n"
            f"📍 Ubicación: {filename}\n\n"
            f"✅ MEJORAS APLICADAS:\n"
            f"   • Encabezados como marca de agua\n"
            f"   • Títulos en color negro\n"
            f"   • Niveles de esquema correctos\n"
            f"   • Formato de citas mejorado\n"
            f"   • Referencias APA optimizadas\n\n"
            f"📋 PARA COMPLETAR EN WORD:\n"
            f"   • Referencias > Tabla de contenido > Automática\n"
            f"   • El índice detectará todos los niveles\n\n"
            f"🚀 ¡Tu proyecto está listo con calidad profesional!"
        )
        messagebox.showinfo("🎉 ¡Éxito Total!", 
            f"Documento generado con todas las mejoras:\n{filename}\n\n"
            f"Características implementadas:\n"
            f"• Encabezados como marca de agua\n"
            f"• Títulos en negro (no azul)\n"
            f"• Niveles de esquema funcionales\n"
            f"• Sistema de citas optimizado\n"
            f"• Formato profesional completo")

